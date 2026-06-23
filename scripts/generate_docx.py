import os
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_document():
    doc = docx.Document()

    # Configuração de margens (1 polegada nas bordas)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilos de cores
    PRIMARY_COLOR = RGBColor(13, 33, 55)   # #0D2137
    SECONDARY_COLOR = RGBColor(21, 101, 192) # #1565C0
    TEXT_COLOR = RGBColor(33, 33, 33)      # #212121

    # Título Principal
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Mapa de Permissões e Controle de Acesso (RBAC) — Relm Care+")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(24)

    # Introdução
    p_intro = doc.add_paragraph()
    intro_run = p_intro.add_run(
        "Este documento apresenta a estrutura completa de papéis, a hierarquia de controle de "
        "acesso baseado em funções (RBAC) e o mapeamento de endpoints do sistema Relm Care+."
    )
    intro_run.font.name = 'Arial'
    intro_run.font.size = Pt(11)
    intro_run.font.color.rgb = TEXT_COLOR
    p_intro.paragraph_format.space_after = Pt(18)

    # Seção 1: Hierarquia de Papéis
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Hierarquia de Papéis e Acesso")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(14)
    h1_run.font.bold = True
    h1_run.font.color.rgb = SECONDARY_COLOR
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    p_h1 = doc.add_paragraph()
    p_h1_run = p_h1.add_run(
        "A estrutura de permissões do sistema divide-se em 6 perfis principais. Os perfis administrativos "
        "da equipe interna (ADMIN_RELM, GERENTE_RELM, SUPORTE_RELM) operam em regime de herança vertical de acessos, "
        "enquanto os perfis parceiros e externos (LOJA, DISTRIBUIDOR, CUSTOMER) possuem acessos isolados de acordo com o escopo."
    )
    p_h1_run.font.name = 'Arial'
    p_h1_run.font.size = Pt(11)
    p_h1_run.font.color.rgb = TEXT_COLOR
    p_h1.paragraph_format.space_after = Pt(18)

    # Seção 2: Mapeamento Detalhado por Perfil
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Mapeamento Detalhado por Perfil")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(14)
    h2_run.font.bold = True
    h2_run.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    profiles = [
        ("Administrador (ADMIN_RELM)", [
            "Escopo: Acesso irrestrito a todos os dados e configurações do sistema.",
            "Permissões Exclusivas: Visualização completa de logs de auditoria (GET /audit-logs), gerenciamento de usuários administrativos (CRUD de User) e deleção de lojas (DELETE /stores/:id), eventos, benefícios e banners."
        ]),
        ("Gerente (GERENTE_RELM)", [
            "Escopo: Gestão operacional diária de revendedores, garantias e campanhas de vantagens.",
            "Permissões Principais: Aprovação e rejeição de solicitações de garantias, aprovação de cotações de seguros, criação e atualização de lojas, eventos e benefícios.",
            "Restrições: Não possui permissão para deletar lojas, acessar logs de auditoria ou gerenciar equipe administrativa."
        ]),
        ("Suporte (SUPORTE_RELM)", [
            "Escopo: Atendimento técnico de primeiro nível e triagem de garantias.",
            "Permissões Principais: Transição de status de garantias na FSM (ex: de RECEBIDO para EM_ANALISE ou AGUARDANDO_CLIENTE), visualização de clientes, seguros, eventos e banners.",
            "Restrições: Não possui permissão de aprovação final (aprovar/rejeitar garantia ou seguro), criação ou deleção de registros."
        ]),
        ("Loja / Lojista (LOJA / STORE)", [
            "Escopo: Acesso restrito para revendedores autorizados da Relm Bikes.",
            "Permissões Principais: Visualização de garantias, clientes e seguros vinculados especificamente à sua própria loja. Registro de garantias e solicitação de cotações de seguros. Gerenciamento de seus próprios colaboradores no portal da loja.",
            "Restrições: Isolamento total contra visualização de dados de outras lojas parceiras."
        ]),
        ("Distribuidor (DISTRIBUIDOR)", [
            "Escopo: Acesso de acompanhamento comercial para distribuidores parceiros.",
            "Permissões Principais: Listar e visualizar dados gerais de clientes e lojas para relatórios de vendas.",
            "Restrições de Privacidade (LGPD): Todos os dados sensíveis do cliente (como CPF e Telefone) são retornados com máscaras de privacidade (ex: 123.***.**-01)."
        ]),
        ("Cliente (CLIENTE / CUSTOMER)", [
            "Escopo: Área exclusiva de clientes finais (auto-atendimento).",
            "Permissões Principais: Visualização do histórico de suas próprias garantias, cotações de seguros contratadas e eventos inscritos. Auto-cadastro integrado.",
            "Restrições: Totalmente isolado do painel e das APIs operacionais e administrativas."
        ])
    ]

    for title, items in profiles:
        sh = doc.add_paragraph()
        sh_run = sh.add_run(f"2.{profiles.index((title, items)) + 1}. {title}")
        sh_run.font.name = 'Arial'
        sh_run.font.size = Pt(12)
        sh_run.font.bold = True
        sh_run.font.color.rgb = PRIMARY_COLOR
        sh.paragraph_format.space_before = Pt(8)
        sh.paragraph_format.space_after = Pt(4)

        for item in items:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet_run = bullet.add_run(item)
            bullet_run.font.name = 'Arial'
            bullet_run.font.size = Pt(11)
            bullet_run.font.color.rgb = TEXT_COLOR
            bullet.paragraph_format.space_after = Pt(2)
        
        # Add space after each profile block
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Seção 3: Matriz Geral de Endpoints
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Matriz Geral de Endpoints (RBAC)")
    h3_run.font.name = 'Arial'
    h3_run.font.size = Pt(14)
    h3_run.font.bold = True
    h3_run.font.color.rgb = SECONDARY_COLOR
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    p_h3 = doc.add_paragraph()
    p_h3_run = p_h3.add_run(
        "A tabela abaixo resume a matriz de permissões para os principais endpoints expostos no backend:"
    )
    p_h3_run.font.name = 'Arial'
    p_h3_run.font.size = Pt(11)
    p_h3_run.font.color.rgb = TEXT_COLOR
    p_h3.paragraph_format.space_after = Pt(12)

    # Matriz de dados
    headers = ["Funcionalidade", "Endpoint", "ADMIN", "GERENTE", "SUPORTE", "LOJA", "DIST.", "CLIE."]
    data = [
        ["Garantia — Registrar (Público)", "POST /public/warranty", "SIM", "SIM", "SIM", "SIM", "SIM", "SIM"],
        ["Garantia — Listar", "GET /warranty/claims", "SIM", "SIM", "SIM", "NAO", "NAO", "NAO"],
        ["Garantia — Alterar Status (FSM)", "PATCH /warranty/claims/:id/status", "SIM", "SIM", "SIM", "NAO", "NAO", "NAO"],
        ["Garantia — Aprovar / Rejeitar", "POST /warranty/claims/:id/approve|reject", "SIM", "SIM", "NAO", "NAO", "NAO", "NAO"],
        ["Garantia — Validar Comprovante", "GET /public/warranty/validate/:token", "SIM", "SIM", "SIM", "SIM", "SIM", "SIM"],
        ["Garantia — Minhas Garantias", "GET /customer-portal/warranties", "NAO", "NAO", "NAO", "NAO", "NAO", "SIM"],
        ["Lojas — Criar / Atualizar", "POST / PATCH /stores", "SIM", "SIM", "NAO", "NAO", "NAO", "NAO"],
        ["Lojas — Listar", "GET /stores", "SIM", "SIM", "SIM", "NAO", "SIM", "NAO"],
        ["Lojas — Excluir", "DELETE /stores/:id", "SIM", "NAO", "NAO", "NAO", "NAO", "NAO"],
        ["Clientes — Listar / Detalhes", "GET /customers", "SIM", "SIM", "SIM", "Masc.", "Masc.", "NAO"],
        ["Clientes — Remover", "DELETE /customers/:id", "SIM", "NAO", "NAO", "NAO", "NAO", "NAO"],
        ["Seguros — Listar / Detalhes", "GET /insurance/quotes", "SIM", "SIM", "SIM", "SIM", "NAO", "NAO"],
        ["Seguros — Aprovar / Rejeitar", "PATCH /insurance/quotes/:id/approve", "SIM", "SIM", "NAO", "NAO", "NAO", "NAO"],
        ["Seguros — Minhas Cotações", "GET /customer-portal/insurance-quotes", "NAO", "NAO", "NAO", "NAO", "NAO", "SIM"],
        ["Eventos — Criar / Editar", "POST / PATCH /events", "SIM", "SIM", "NAO", "NAO", "NAO", "NAO"],
        ["Eventos — Inscrições", "GET /events/:id/registrations", "SIM", "SIM", "SIM", "NAO", "NAO", "NAO"],
        ["Benefícios — Criar / Editar", "POST / PATCH /benefits", "SIM", "SIM", "NAO", "NAO", "NAO", "NAO"],
        ["Auditoria — Visualizar Logs", "GET /audit-logs", "SIM", "NAO", "NAO", "NAO", "NAO", "NAO"],
        ["Usuários Admin — Gerenciar", "* /admin-users", "SIM", "NAO", "NAO", "NAO", "NAO", "NAO"]
    ]

    # Criando a tabela no documento
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'

    # Estilizando o cabeçalho da tabela
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "0D2137")
        # Fonte branca e bold
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Populando as linhas da tabela
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        # Linhas zebradas (fundo cinza claro para linhas pares)
        bg_color = "F5F5F5" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            set_cell_background(row_cells[col_idx], bg_color)
            p = row_cells[col_idx].paragraphs[0]
            # Formatações específicas de alinhamento
            if col_idx >= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
                    if cell_value == "SIM":
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(56, 142, 60) # Verde
                    elif cell_value == "NAO":
                        run.font.color.rgb = RGBColor(198, 40, 40) # Vermelho
                    elif cell_value == "Masc.":
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(21, 101, 192) # Azul
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
                    run.font.color.rgb = TEXT_COLOR

    # Legenda
    p_legenda = doc.add_paragraph()
    p_legenda.paragraph_format.space_before = Pt(12)
    legenda_run = p_legenda.add_run(
        "Legenda: SIM = Acesso permitido; NAO = Acesso negado; "
        "Masc. = Acesso permitido com dados de CPF e Telefone mascarados (LGPD)."
    )
    legenda_run.font.name = 'Arial'
    legenda_run.font.size = Pt(9)
    legenda_run.font.italic = True
    legenda_run.font.color.rgb = RGBColor(117, 117, 117)

    # Salva o arquivo no workspace principal
    dest_path = r"c:\Users\BOSS\Desktop\Relm\Relm-Care\permissions_map.docx"
    doc.save(dest_path)
    print(f"Documento de permissões salvo com sucesso em: {dest_path}")

if __name__ == '__main__':
    create_document()
